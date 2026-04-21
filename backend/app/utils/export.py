"""简历导出工具 - Markdown/结构化数据 → PDF / DOCX"""

import os
import uuid

import markdown as md

from app.config.settings import settings

# ── 通用 CSS ──────────────────────────────────────

RESUME_CSS = """
body {
    font-family: "Noto Sans SC", "Microsoft YaHei", "Helvetica Neue", Arial, sans-serif;
    font-size: 11pt;
    line-height: 1.6;
    color: #333;
    max-width: 800px;
    margin: 0 auto;
    padding: 40px;
}
h1 {
    font-size: 22pt;
    color: #1a1a1a;
    border-bottom: 2px solid #333;
    padding-bottom: 8px;
    margin-bottom: 4px;
}
h2 {
    font-size: 14pt;
    color: #2c3e50;
    border-bottom: 1px solid #ddd;
    padding-bottom: 4px;
    margin-top: 16px;
}
h3 { font-size: 12pt; color: #34495e; }
hr { border: none; border-top: 1px solid #ddd; margin: 8px 0; }
ul { padding-left: 20px; }
li { margin-bottom: 4px; }
p { margin: 6px 0; }
strong { color: #1a1a1a; }
"""

# ── 模板 CSS 定义 ──────────────────────────────────────

TEMPLATE_CSS = {
    "professional": """
body { font-family: "Noto Sans SC", "Georgia", serif; font-size: 11pt; color: #333; margin: 0; padding: 0; }
.resume { max-width: 800px; margin: 0 auto; padding: 40px; }
.header { background: #1e3a5f; color: white; padding: 24px 32px; margin: -40px -40px 24px; }
.header h1 { font-size: 24pt; margin: 0 0 4px; border: none; color: white; }
.header .title { font-size: 13pt; opacity: 0.9; }
.header .contact { font-size: 10pt; opacity: 0.8; margin-top: 8px; }
.header .contact span { margin-right: 16px; }
.section { margin-bottom: 16px; }
.section h2 { font-size: 13pt; color: #1e3a5f; border-bottom: 2px solid #1e3a5f; padding-bottom: 4px; margin: 0 0 10px; text-transform: uppercase; letter-spacing: 1px; }
.entry { margin-bottom: 12px; }
.entry-header { display: flex; justify-content: space-between; align-items: baseline; margin-bottom: 2px; }
.entry-header h3 { font-size: 11pt; color: #1a1a1a; margin: 0; }
.entry-header .date { font-size: 10pt; color: #666; white-space: nowrap; }
.entry-sub { font-size: 10pt; color: #555; margin-bottom: 4px; }
ul { padding-left: 18px; margin: 4px 0; }
li { font-size: 10.5pt; margin-bottom: 3px; line-height: 1.5; }
.skills-grid { display: flex; flex-wrap: wrap; gap: 8px 24px; }
.skill-group { font-size: 10.5pt; }
.skill-group strong { color: #1e3a5f; }
.summary { font-size: 10.5pt; line-height: 1.6; color: #444; font-style: italic; }
""",
    "minimalist": """
body { font-family: "Noto Sans SC", "Helvetica Neue", Arial, sans-serif; font-size: 10.5pt; color: #222; margin: 0; padding: 0; }
.resume { max-width: 720px; margin: 0 auto; padding: 48px 40px; }
.header { text-align: center; margin-bottom: 24px; border-bottom: 1px solid #ddd; padding-bottom: 16px; }
.header h1 { font-size: 28pt; font-weight: 300; letter-spacing: 4px; margin: 0 0 6px; border: none; color: #111; }
.header .title { font-size: 11pt; color: #666; letter-spacing: 2px; }
.header .contact { font-size: 9.5pt; color: #888; margin-top: 10px; }
.header .contact span { margin: 0 8px; }
.section { margin-bottom: 18px; }
.section h2 { font-size: 10pt; color: #999; text-transform: uppercase; letter-spacing: 3px; border-bottom: none; margin: 0 0 10px; font-weight: 400; }
.entry { margin-bottom: 12px; }
.entry-header { display: flex; justify-content: space-between; align-items: baseline; }
.entry-header h3 { font-size: 10.5pt; font-weight: 600; margin: 0; }
.entry-header .date { font-size: 9.5pt; color: #999; }
.entry-sub { font-size: 9.5pt; color: #666; }
ul { padding-left: 16px; margin: 4px 0; }
li { font-size: 10pt; margin-bottom: 2px; color: #444; }
.skills-grid { display: flex; flex-wrap: wrap; gap: 6px 20px; font-size: 10pt; }
.summary { font-size: 10pt; color: #555; line-height: 1.7; text-align: center; }
""",
    "creative": """
body { font-family: "Noto Sans SC", "Helvetica Neue", sans-serif; font-size: 10.5pt; color: #333; margin: 0; padding: 0; }
.resume { display: flex; min-height: 100vh; }
.sidebar { width: 240px; background: linear-gradient(135deg, #6366f1, #8b5cf6); color: white; padding: 32px 20px; flex-shrink: 0; }
.sidebar h1 { font-size: 20pt; margin: 0 0 4px; border: none; color: white; }
.sidebar .title { font-size: 11pt; opacity: 0.85; margin-bottom: 20px; }
.sidebar .contact { font-size: 9.5pt; opacity: 0.8; }
.sidebar .contact div { margin-bottom: 6px; }
.sidebar .section h2 { font-size: 10pt; color: rgba(255,255,255,0.7); text-transform: uppercase; letter-spacing: 2px; border-bottom: 1px solid rgba(255,255,255,0.2); padding-bottom: 4px; margin: 16px 0 8px; }
.sidebar .skill-item { display: flex; justify-content: space-between; font-size: 9.5pt; margin-bottom: 4px; }
.sidebar .skill-bar { height: 4px; background: rgba(255,255,255,0.2); border-radius: 2px; margin-top: 2px; }
.sidebar .skill-bar-fill { height: 100%; background: white; border-radius: 2px; }
.main { flex: 1; padding: 32px; }
.main .section { margin-bottom: 18px; }
.main .section h2 { font-size: 12pt; color: #6366f1; border-bottom: 2px solid #e0e0ff; padding-bottom: 4px; margin: 0 0 10px; }
.entry { margin-bottom: 14px; padding-left: 12px; border-left: 3px solid #e0e0ff; }
.entry-header h3 { font-size: 11pt; margin: 0; color: #1a1a1a; }
.entry-header .date { font-size: 9.5pt; color: #8b5cf6; }
.entry-sub { font-size: 9.5pt; color: #666; }
ul { padding-left: 16px; margin: 4px 0; }
li { font-size: 10pt; margin-bottom: 2px; }
.summary { font-size: 10.5pt; color: #444; line-height: 1.6; }
""",
    "tech": """
body { font-family: "Noto Sans SC", "Fira Code", "JetBrains Mono", monospace; font-size: 10.5pt; color: #e0e0e0; background: #0f172a; margin: 0; padding: 0; }
.resume { max-width: 800px; margin: 0 auto; padding: 40px; }
.header { background: #1e293b; border: 1px solid #334155; border-radius: 8px; padding: 24px; margin-bottom: 20px; }
.header h1 { font-size: 22pt; color: #22d3ee; margin: 0 0 4px; border: none; font-family: "Fira Code", monospace; }
.header .title { font-size: 11pt; color: #94a3b8; }
.header .contact { font-size: 9.5pt; color: #64748b; margin-top: 8px; }
.header .contact span { margin-right: 16px; }
.header .contact a { color: #22d3ee; text-decoration: none; }
.section { margin-bottom: 16px; }
.section h2 { font-size: 11pt; color: #22d3ee; border-bottom: 1px solid #334155; padding-bottom: 4px; margin: 0 0 10px; font-family: "Fira Code", monospace; }
.section h2::before { content: "// "; color: #475569; }
.entry { margin-bottom: 12px; background: #1e293b; border-radius: 6px; padding: 12px 16px; border: 1px solid #334155; }
.entry-header h3 { font-size: 10.5pt; color: #f1f5f9; margin: 0; }
.entry-header .date { font-size: 9.5pt; color: #22d3ee; font-family: "Fira Code", monospace; }
.entry-sub { font-size: 9.5pt; color: #94a3b8; }
ul { padding-left: 16px; margin: 4px 0; }
li { font-size: 10pt; margin-bottom: 2px; color: #cbd5e1; }
.skills-grid { display: flex; flex-wrap: wrap; gap: 6px; }
.skill-tag { background: #1e293b; border: 1px solid #22d3ee; color: #22d3ee; padding: 2px 10px; border-radius: 4px; font-size: 9.5pt; font-family: "Fira Code", monospace; }
.summary { font-size: 10pt; color: #94a3b8; line-height: 1.6; }
@media print { body { background: white; color: #333; } .resume { padding: 20px; } .header, .entry { background: #f8f9fa; border-color: #ddd; } .section h2, .header h1, .entry-header .date, .skill-tag { color: #0891b2; } .entry-header h3 { color: #1a1a1a; } li, .summary, .entry-sub { color: #444; } .header .contact { color: #666; } .skill-tag { border-color: #0891b2; } }
""",
    "academic": """
body { font-family: "Noto Serif SC", "Times New Roman", "Georgia", serif; font-size: 11pt; color: #1a1a1a; margin: 0; padding: 0; }
.resume { max-width: 760px; margin: 0 auto; padding: 48px 40px; }
.header { text-align: center; margin-bottom: 20px; }
.header h1 { font-size: 20pt; font-weight: 400; letter-spacing: 2px; margin: 0 0 6px; border: none; }
.header .title { font-size: 11pt; color: #555; }
.header .contact { font-size: 10pt; color: #666; margin-top: 8px; }
.header .contact span { margin: 0 6px; }
.section { margin-bottom: 14px; }
.section h2 { font-size: 11pt; font-variant: small-caps; letter-spacing: 2px; border-bottom: 1px solid #333; padding-bottom: 2px; margin: 0 0 8px; font-weight: 400; }
.entry { margin-bottom: 10px; }
.entry-header { display: flex; justify-content: space-between; align-items: baseline; }
.entry-header h3 { font-size: 11pt; font-weight: 600; margin: 0; font-style: italic; }
.entry-header .date { font-size: 10pt; color: #555; }
.entry-sub { font-size: 10pt; color: #444; margin-left: 16px; }
ul { padding-left: 20px; margin: 4px 0; }
li { font-size: 10.5pt; margin-bottom: 2px; }
.summary { font-size: 10.5pt; line-height: 1.7; text-indent: 2em; }
.skills-grid { font-size: 10.5pt; }
.skill-group { margin-bottom: 4px; }
""",
    "executive": """
body { font-family: "Noto Sans SC", "Garamond", "Georgia", serif; font-size: 11pt; color: #2c2c2c; margin: 0; padding: 0; }
.resume { max-width: 800px; margin: 0 auto; padding: 48px 44px; }
.header { border-bottom: 2px solid #b8860b; padding-bottom: 16px; margin-bottom: 24px; }
.header h1 { font-size: 26pt; font-weight: 300; letter-spacing: 3px; text-transform: uppercase; margin: 0 0 4px; border: none; color: #1a1a1a; }
.header .title { font-size: 12pt; color: #b8860b; letter-spacing: 1px; }
.header .contact { font-size: 10pt; color: #777; margin-top: 10px; }
.header .contact span { margin-right: 16px; }
.section { margin-bottom: 20px; }
.section h2 { font-size: 11pt; color: #b8860b; text-transform: uppercase; letter-spacing: 2px; border-bottom: 1px solid #e0d5b8; padding-bottom: 4px; margin: 0 0 12px; font-weight: 400; }
.entry { margin-bottom: 14px; }
.entry-header { display: flex; justify-content: space-between; align-items: baseline; }
.entry-header h3 { font-size: 11.5pt; font-weight: 600; margin: 0; color: #1a1a1a; }
.entry-header .date { font-size: 10pt; color: #888; }
.entry-sub { font-size: 10pt; color: #666; font-style: italic; }
ul { padding-left: 18px; margin: 6px 0; }
li { font-size: 10.5pt; margin-bottom: 4px; line-height: 1.6; }
.summary { font-size: 11pt; line-height: 1.7; color: #333; border-left: 3px solid #b8860b; padding-left: 16px; font-style: italic; }
.skills-grid { display: flex; flex-wrap: wrap; gap: 8px 24px; font-size: 10.5pt; }
.skill-group strong { color: #b8860b; }
""",
}


def _render_resume_html(data: dict, template_id: str) -> str:
    """将结构化数据渲染为 HTML"""
    personal = data.get("personal", {})
    is_creative = template_id == "creative"

    # Contact info
    contact_parts = []
    for key in ["phone", "email", "location", "website", "linkedin", "github"]:
        val = personal.get(key)
        if val:
            contact_parts.append(f'<span>{val}</span>')
    contact_html = " ".join(contact_parts)

    # Sections
    sections_html = ""

    # Summary
    if data.get("summary"):
        sections_html += f'<div class="section"><h2>个人简介</h2><p class="summary">{data["summary"]}</p></div>'

    # Work Experience
    if data.get("work_experience"):
        entries = ""
        for exp in data["work_experience"]:
            hl = "".join(f"<li>{h}</li>" for h in exp.get("highlights", []))
            entries += f"""<div class="entry">
                <div class="entry-header"><h3>{exp['title']}</h3><span class="date">{exp.get('start_date','')} - {exp.get('end_date','')}</span></div>
                <div class="entry-sub">{exp['company']}{(' · ' + exp['location']) if exp.get('location') else ''}</div>
                <ul>{hl}</ul></div>"""
        sections_html += f'<div class="section"><h2>工作经历</h2>{entries}</div>'

    # Education
    if data.get("education"):
        entries = ""
        for edu in data["education"]:
            hl = "".join(f"<li>{h}</li>" for h in edu.get("highlights", []))
            gpa_str = f" | GPA: {edu['gpa']}" if edu.get("gpa") else ""
            entries += f"""<div class="entry">
                <div class="entry-header"><h3>{edu['degree']} - {edu['field']}</h3><span class="date">{edu.get('start_date','')} - {edu.get('end_date','')}</span></div>
                <div class="entry-sub">{edu['institution']}{gpa_str}</div>
                {'<ul>' + hl + '</ul>' if hl else ''}</div>"""
        sections_html += f'<div class="section"><h2>教育背景</h2>{entries}</div>'

    # Skills
    if data.get("skills"):
        if template_id == "tech":
            tags = ""
            for sg in data["skills"]:
                for item in sg.get("items", []):
                    tags += f'<span class="skill-tag">{item}</span>'
            sections_html += f'<div class="section"><h2>专业技能</h2><div class="skills-grid">{tags}</div></div>'
        else:
            items = ""
            for sg in data["skills"]:
                items += f'<div class="skill-group"><strong>{sg["category"]}</strong>: {", ".join(sg.get("items", []))}</div>'
            sections_html += f'<div class="section"><h2>专业技能</h2><div class="skills-grid">{items}</div></div>'

    # Projects
    if data.get("projects"):
        entries = ""
        for proj in data["projects"]:
            hl = "".join(f"<li>{h}</li>" for h in proj.get("highlights", []))
            tech = ""
            if proj.get("tech_stack"):
                if template_id == "tech":
                    tech = '<div style="margin-top:4px">' + "".join(f'<span class="skill-tag">{t}</span>' for t in proj["tech_stack"]) + "</div>"
                else:
                    tech = f'<div class="entry-sub">技术栈: {", ".join(proj["tech_stack"])}</div>'
            date_str = ""
            if proj.get("start_date"):
                date_str = f'{proj["start_date"]} - {proj.get("end_date", "")}'
            entries += f"""<div class="entry">
                <div class="entry-header"><h3>{proj['name']}{(' - ' + proj['role']) if proj.get('role') else ''}</h3><span class="date">{date_str}</span></div>
                <p class="entry-sub">{proj.get('description','')}</p>
                {tech}<ul>{hl}</ul></div>"""
        sections_html += f'<div class="section"><h2>项目经验</h2>{entries}</div>'

    # Certifications
    if data.get("certifications"):
        items = ""
        for cert in data["certifications"]:
            line = cert["name"]
            if cert.get("issuer"):
                line += f' ({cert["issuer"]})'
            if cert.get("date"):
                line += f' - {cert["date"]}'
            items += f"<li>{line}</li>"
        sections_html += f'<div class="section"><h2>证书与奖项</h2><ul>{items}</ul></div>'

    # Custom Sections
    for sec in data.get("custom_sections", []):
        sections_html += f'<div class="section"><h2>{sec["title"]}</h2><p>{sec["content"]}</p></div>'

    # Assemble based on template layout
    if is_creative:
        # Two-column layout: sidebar + main
        sidebar_sections = ""
        # Put contact and skills in sidebar
        sidebar_contact = "".join(f'<div>{p}</div>' for p in [
            personal.get("phone", ""),
            personal.get("email", ""),
            personal.get("location", ""),
            personal.get("website", ""),
            personal.get("github", ""),
        ] if p)

        sidebar_skills = ""
        if data.get("skills"):
            skill_items = ""
            for sg in data["skills"]:
                for item in sg.get("items", []):
                    skill_items += f'<div class="skill-item"><span>{item}</span></div>'
            sidebar_skills = f'<div class="section"><h2>技能</h2>{skill_items}</div>'

        return f"""<div class="resume">
            <div class="sidebar">
                <h1>{personal.get('name','')}</h1>
                <div class="title">{personal.get('title','')}</div>
                <div class="contact">{sidebar_contact}</div>
                {sidebar_skills}
            </div>
            <div class="main">{sections_html}</div>
        </div>"""
    else:
        return f"""<div class="resume">
            <div class="header">
                <h1>{personal.get('name','')}</h1>
                <div class="title">{personal.get('title','')}</div>
                <div class="contact">{contact_html}</div>
            </div>
            {sections_html}
        </div>"""


# ── Markdown 导出 (旧版兼容) ──────────────────────────────

def markdown_to_pdf(content: str) -> str:
    """Markdown → PDF, 返回文件路径 (需要 weasyprint)"""
    try:
        from weasyprint import HTML
    except ImportError:
        raise RuntimeError("PDF 导出需要安装 weasyprint 及其系统依赖")

    html_body = md.markdown(content, extensions=["tables", "fenced_code", "nl2br"])
    full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>{RESUME_CSS}</style></head>
<body>{html_body}</body></html>"""

    filename = f"resume_{uuid.uuid4().hex[:8]}.pdf"
    output_path = os.path.join(settings.upload_dir, filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    HTML(string=full_html).write_pdf(output_path)
    return output_path


def markdown_to_docx(content: str) -> str:
    """Markdown → DOCX, 返回文件路径"""
    import re
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)

    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            p = doc.add_paragraph()
            parts = re.split(r"(\*\*[^*]+\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)

    filename = f"resume_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(settings.upload_dir, filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path


# ── 结构化数据导出 ──────────────────────────────────────

def structured_to_pdf(data: dict, template_id: str = "professional") -> str:
    """结构化简历数据 + 模板 → PDF (需要 weasyprint)"""
    try:
        from weasyprint import HTML
    except ImportError:
        raise RuntimeError("PDF 导出需要安装 weasyprint 及其系统依赖")

    css = TEMPLATE_CSS.get(template_id, TEMPLATE_CSS["professional"])
    body_html = _render_resume_html(data, template_id)

    full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>{css}</style></head>
<body>{body_html}</body></html>"""

    filename = f"resume_{uuid.uuid4().hex[:8]}.pdf"
    output_path = os.path.join(settings.upload_dir, filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)

    HTML(string=full_html).write_pdf(output_path)
    return output_path


def structured_to_docx(data: dict, template_id: str = "professional") -> str:
    """结构化简历数据 → DOCX"""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Microsoft YaHei"
    font.size = Pt(11)

    personal = data.get("personal", {})

    # Name
    p = doc.add_heading(personal.get("name", ""), level=1)
    if personal.get("title"):
        p = doc.add_paragraph(personal["title"])
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact
    contact_items = [personal.get(k) for k in ["phone", "email", "location", "website"] if personal.get(k)]
    if contact_items:
        p = doc.add_paragraph(" | ".join(contact_items))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Summary
    if data.get("summary"):
        doc.add_heading("个人简介", level=2)
        doc.add_paragraph(data["summary"])

    # Work Experience
    if data.get("work_experience"):
        doc.add_heading("工作经历", level=2)
        for exp in data["work_experience"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{exp['title']}  ")
            run.bold = True
            run.font.size = Pt(11)
            p.add_run(f"@ {exp['company']}").font.size = Pt(10)
            p = doc.add_paragraph(f"{exp.get('start_date', '')} - {exp.get('end_date', '')}")
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            for h in exp.get("highlights", []):
                doc.add_paragraph(h, style="List Bullet")

    # Education
    if data.get("education"):
        doc.add_heading("教育背景", level=2)
        for edu in data["education"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu['degree']} - {edu['field']}  ")
            run.bold = True
            p.add_run(f"@ {edu['institution']}")
            p = doc.add_paragraph(f"{edu.get('start_date', '')} - {edu.get('end_date', '')}")
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Skills
    if data.get("skills"):
        doc.add_heading("专业技能", level=2)
        for sg in data["skills"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{sg['category']}: ")
            run.bold = True
            p.add_run(", ".join(sg.get("items", [])))

    # Projects
    if data.get("projects"):
        doc.add_heading("项目经验", level=2)
        for proj in data["projects"]:
            p = doc.add_paragraph()
            run = p.add_run(proj["name"])
            run.bold = True
            if proj.get("description"):
                doc.add_paragraph(proj["description"])
            for h in proj.get("highlights", []):
                doc.add_paragraph(h, style="List Bullet")

    # Certifications
    if data.get("certifications"):
        doc.add_heading("证书与奖项", level=2)
        for cert in data["certifications"]:
            line = cert["name"]
            if cert.get("issuer"):
                line += f" ({cert['issuer']})"
            if cert.get("date"):
                line += f" - {cert['date']}"
            doc.add_paragraph(line, style="List Bullet")

    filename = f"resume_{uuid.uuid4().hex[:8]}.docx"
    output_path = os.path.join(settings.upload_dir, filename)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    return output_path
